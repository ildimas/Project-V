from docx import Document
import os
from docxcompose.composer import Composer

class DocxPretier:
    def __init__(self, original_filepath : str, changed_filepath : str):
        self.original_filepath = original_filepath
        self.changed_filepath = changed_filepath
        self.process()
        
    def process(self):
        try:
            doc_path = self.original_filepath
            doc = Document(doc_path)

            if doc.paragraphs:
                first_paragraph = doc.paragraphs[0]
                for run in first_paragraph.runs:
                    run.clear()

            for section in doc.sections:
                footer = section.footer
                footer.paragraphs[0].clear()
                if section.start_type not in [None, 'continuous']:
                    section.first_page_footer.paragraphs[0].clear()
                if section.even_page_footer is not None:
                    section.even_page_footer.paragraphs[0].clear()
            
            modified_doc_path = self.changed_filepath
            doc.save(modified_doc_path)
            modified_doc_path
            # os.remove(self.original_filepath)
            print("File pretier finished his work!")
        except Exception as e:
            print(f"Error ocured in module 1 {__name__} exception : {e}")
        
        
class DocxCompressor:
    def __init__(self, files_list : list, combined_file : str):
         self.files_list = files_list 
         self.combined_file = combined_file
         self.compressor()
         
    def compressor(self):
        try:
            master = Document(self.files_list[0])
            composer = Composer(master)

            for file in self.files_list[1:]:
                print("docx_pret:",file)
                doc_to_merge = Document(file)
                composer.append(doc_to_merge)

            composer.save(self.combined_file)
            print("Compressor finished !")
        except Exception as e:
            print(f"Error occurred in module 2 {__name__} exception: {e}")
         

                
if __name__ == "__main__": 
    # DocxPretier("docs/results/Output_1.docx", "docs/results/New_Output_1.docx")
    DocxCompressor(["docs/results/mainsample2.docx", "docs/results/mainsample.docx"], "docs/results/файл.docx")