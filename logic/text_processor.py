import docx
import fitz 
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx import Document
import os




class TextProcessor():
    def __init__(self):
        pass

    def extract_text_from_docx(self, filepath):
        try:
            doc = docx.Document(filepath)
            full_text = ""
            for para in doc.paragraphs:
                full_text += para.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
                    full_text += "\n"
                full_text += "\n"
            
            return full_text
        except Exception as e:
            print(f"Exception in module {__name__} function extract_text_from_docx: {e}")
    
    def extract_text_from_txt(self, filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                # return self.insp.normalize_and_split(file.read())
                return file.read()
        except Exception as e:
            print(f"Exception in module {__name__} function extract_from_txt")
        
    def extract_text_from_pdf_with_fitz(self, filepath):
        try:
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            return text
        except Exception as e:
                print(f"Exception in module {__name__} function extract_from_pdf")

if __name__ == "__main__":
    x = TextProcessor()
    # print(x.extract_text_from_pdf_with_pypdf2(r"C:\Users\spoyi\OneDrive\Рабочий стол\pod trade stuff\Compare mechanism\File_compare_server_build\logic\samples\Тяжелый договор 1.pdf"))
    # print("****************")
    # # print(x.extract_text_from_txt("samples/example.txt"))
    # # print("****************")
    # # print(x.extract_text_from_pdf_with_fitz("samples/mainsample.pdf"))
    x.split_docx("samples\Тяжелый договор 2.docx", "docs/after/docx")
